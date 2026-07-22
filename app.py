import os
import re
import json
import hashlib
import base64
import smtplib
import ipaddress
import socket
from datetime import datetime, timedelta
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from functools import wraps
from urllib.parse import urlparse

from flask import (
    Flask, render_template, request, jsonify, redirect,
    url_for, flash, Response, stream_with_context
)
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from flask_bcrypt import Bcrypt
from flask_wtf.csrf import CSRFProtect
from flask_session import Session
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from cryptography.fernet import Fernet
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired
from apscheduler.schedulers.background import BackgroundScheduler

import stripe
import requests as req_lib
from bs4 import BeautifulSoup
import openpyxl
import csv
import io

from models import db, User, Campaign, CampaignRecipient, ScheduledCampaign

app = Flask(__name__)

# ── Config ────────────────────────────────────────────────────────────────────
_secret = os.environ.get("SECRET_KEY") or os.environ.get("SESSION_SECRET")
if not _secret:
    raise RuntimeError(
        "SECRET_KEY environment variable is not set. "
        "Set it to a strong random value before starting the server."
    )
app.config["SECRET_KEY"] = _secret
app.config["SQLALCHEMY_DATABASE_URI"] = os.environ.get(
    "DATABASE_URL", "sqlite:///rushmail.db"
).replace("postgres://", "postgresql://")
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["WTF_CSRF_TIME_LIMIT"] = None
app.config["MAX_CONTENT_LENGTH"] = 32 * 1024 * 1024  # 32 MB upload limit

# ── Server-side sessions ──────────────────────────────────────────────────────
app.config["SESSION_TYPE"] = "filesystem"
app.config["SESSION_FILE_DIR"] = os.path.join(os.getcwd(), ".flask_sessions")
app.config["SESSION_FILE_THRESHOLD"] = 500
app.config["SESSION_PERMANENT"] = False

# Allow the session cookie to survive being loaded inside a cross-site iframe
# (e.g. the Replit preview/canvas embeds the app in an iframe on another
# origin). SameSite=None requires Secure=True, which is safe here since the
# app is always served over HTTPS in both dev preview and production.
app.config["SESSION_COOKIE_SAMESITE"] = "None"
app.config["SESSION_COOKIE_SECURE"] = True
app.config["REMEMBER_COOKIE_SAMESITE"] = "None"
app.config["REMEMBER_COOKIE_SECURE"] = True

db.init_app(app)
bcrypt = Bcrypt(app)
csrf = CSRFProtect(app)
Session(app)

login_manager = LoginManager(app)
login_manager.login_view = "login"
login_manager.login_message_category = "error"

limiter = Limiter(
    key_func=get_remote_address,
    app=app,
    default_limits=[],
    storage_uri="memory://",
)

# ── Stripe ────────────────────────────────────────────────────────────────────
STRIPE_SECRET_KEY = os.environ.get("STRIPE_SECRET_KEY", "")
STRIPE_PUBLISHABLE_KEY = os.environ.get("STRIPE_PUBLISHABLE_KEY", "")
STRIPE_WEBHOOK_SECRET = os.environ.get("STRIPE_WEBHOOK_SECRET", "")
STRIPE_PRO_PRICE_ID = os.environ.get("STRIPE_PRO_PRICE_ID", "")
STRIPE_ENABLED = bool(STRIPE_SECRET_KEY)

if STRIPE_ENABLED:
    stripe.api_key = STRIPE_SECRET_KEY

# ── OpenAI ────────────────────────────────────────────────────────────────────
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")

# ── Helpers ───────────────────────────────────────────────────────────────────
EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")


def get_fernet():
    raw = app.config["SECRET_KEY"].encode()
    key = base64.urlsafe_b64encode(hashlib.sha256(raw).digest())
    return Fernet(key)


def encrypt_password(plain: str) -> str:
    return get_fernet().encrypt(plain.encode()).decode()


def decrypt_password(token: str) -> str:
    return get_fernet().decrypt(token.encode()).decode()


def extract_emails(text: str) -> list[str]:
    found = EMAIL_RE.findall(text)
    seen = set()
    result = []
    for e in found:
        e_lower = e.lower()
        if e_lower not in seen:
            seen.add(e_lower)
            result.append(e_lower)
    return result


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


def get_domain():
    domain = os.environ.get("REPLIT_DEV_DOMAIN", "")
    if domain:
        return f"https://{domain}"
    return request.host_url.rstrip("/")


# ── Password policy ───────────────────────────────────────────────────────────
_PASS_RE_UPPER   = re.compile(r"[A-Z]")
_PASS_RE_DIGIT   = re.compile(r"\d")
_PASS_RE_SPECIAL = re.compile(r"[!@#$%^&*()\-_=+\[\]{};:',.<>?/\\|`~]")


def validate_password(password: str) -> list[str]:
    """Return a list of unmet password policy requirements (empty = OK)."""
    errors = []
    if len(password) < 8:
        errors.append("at least 8 characters")
    if not _PASS_RE_UPPER.search(password):
        errors.append("one uppercase letter")
    if not _PASS_RE_DIGIT.search(password):
        errors.append("one number")
    if not _PASS_RE_SPECIAL.search(password):
        errors.append("one special character (!@#$%^&* …)")
    return errors


# ── Auth email tokens ─────────────────────────────────────────────────────────
def _get_serializer(salt: str) -> URLSafeTimedSerializer:
    return URLSafeTimedSerializer(app.config["SECRET_KEY"], salt=salt)


def make_verification_token(user_id: int) -> str:
    return _get_serializer("email-verify").dumps(user_id)


def verify_verification_token(token: str, max_age: int = 86400):
    """Return user_id or None (token invalid / expired after 24 h)."""
    try:
        return _get_serializer("email-verify").loads(token, max_age=max_age)
    except (BadSignature, SignatureExpired):
        return None


def make_reset_token(user_id: int) -> str:
    return _get_serializer("password-reset").dumps(user_id)


def verify_reset_token(token: str, max_age: int = 3600):
    """Return user_id or None (token invalid / expired after 1 h)."""
    try:
        return _get_serializer("password-reset").loads(token, max_age=max_age)
    except (BadSignature, SignatureExpired):
        return None


# ── Auth email sender ─────────────────────────────────────────────────────────
# App-level SMTP for transactional auth emails (verify / reset).
# If these env vars are set, they are tried first so that new users who
# haven't configured their own SMTP can still receive verification emails.
# Falls back to the user's own SMTP settings if app-level SMTP is absent.
_APP_SMTP_HOST = os.environ.get("AUTH_SMTP_HOST", "")
_APP_SMTP_PORT = int(os.environ.get("AUTH_SMTP_PORT", "587") or "587")
_APP_SMTP_USER = os.environ.get("AUTH_SMTP_USER", "")
_APP_SMTP_PASS = os.environ.get("AUTH_SMTP_PASS", "")
_APP_SMTP_FROM = os.environ.get("AUTH_SMTP_FROM", "")
_APP_SMTP_TLS  = os.environ.get("AUTH_SMTP_TLS", "true").lower() != "false"


def _smtp_send(host, port, user, password, use_tls, from_addr, to_addr, msg_str) -> bool:
    try:
        if use_tls:
            server = smtplib.SMTP(host, port, timeout=15)
            server.starttls()
        else:
            server = smtplib.SMTP_SSL(host, port, timeout=15)
        server.login(user, password)
        server.sendmail(from_addr, [to_addr], msg_str)
        server.quit()
        return True
    except Exception:
        return False


def _send_auth_email(user: "User", subject: str, body_html: str) -> bool:
    """Send a transactional email for account verification or password reset.

    Tries, in order:
      1. App-level SMTP (AUTH_SMTP_* env vars) — works even for new users
      2. The user's own SMTP settings — works if already configured
    Returns True on success, False if neither is available or the send fails.
    """
    msg = MIMEMultipart("alternative")
    msg["Subject"] = subject
    msg["To"]      = user.email
    msg.attach(MIMEText(body_html, "html"))

    # 1 — App-level SMTP
    if _APP_SMTP_HOST and _APP_SMTP_USER and _APP_SMTP_PASS:
        from_addr = _APP_SMTP_FROM or _APP_SMTP_USER
        msg["From"] = from_addr
        if _smtp_send(_APP_SMTP_HOST, _APP_SMTP_PORT, _APP_SMTP_USER,
                      _APP_SMTP_PASS, _APP_SMTP_TLS, from_addr, user.email,
                      msg.as_string()):
            return True

    # 2 — User's own SMTP
    if user.smtp_host and user.smtp_pass_enc:
        try:
            smtp_pass = decrypt_password(user.smtp_pass_enc)
            from_addr = user.smtp_from or user.smtp_user
            if "From" not in msg:
                msg["From"] = from_addr
            if _smtp_send(user.smtp_host, user.smtp_port, user.smtp_user,
                          smtp_pass, user.smtp_use_tls, from_addr, user.email,
                          msg.as_string()):
                return True
        except Exception:
            pass

    return False


# ── Subscription guard ────────────────────────────────────────────────────────
def subscription_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not current_user.is_pro:
            flash("An active subscription is required. Subscribe below to get started.", "error")
            return redirect(url_for("pricing"))
        return f(*args, **kwargs)
    return decorated


# ── Auth Routes ───────────────────────────────────────────────────────────────
@app.route("/")
def index():
    if current_user.is_authenticated:
        return redirect(url_for("dashboard"))
    return redirect(url_for("login"))


@app.route("/login", methods=["GET", "POST"])
@limiter.limit("5 per 15 minutes", methods=["POST"], error_message="Too many login attempts. Please wait 15 minutes and try again.")
def login():
    if current_user.is_authenticated:
        return redirect(url_for("dashboard"))
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        user = User.query.filter_by(email=email).first()
        if user and bcrypt.check_password_hash(user.password_hash, password):
            if not user.verified:
                flash(
                    "Please verify your email before signing in. "
                    "<a href=\"" + url_for("resend_verification", email=user.email) + "\">Resend verification email</a>",
                    "error"
                )
                return render_template("login.html")
            login_user(user)
            return redirect(url_for("dashboard"))
        flash("Invalid email or password.", "error")
    return render_template("login.html")


@app.route("/register", methods=["GET", "POST"])
def register():
    if current_user.is_authenticated:
        return redirect(url_for("dashboard"))
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        password2 = request.form.get("password2", "")
        policy_errors = validate_password(password)
        if policy_errors:
            return render_template("register.html", password_errors=policy_errors, email=email)
        elif password != password2:
            flash("Passwords do not match.", "error")
            return render_template("register.html", password_errors=[], email=email)
        elif User.query.filter_by(email=email).first():
            flash("An account with that email already exists.", "error")
        else:
            user = User(
                email=email,
                password_hash=bcrypt.generate_password_hash(password).decode(),
                verified=False,
            )
            db.session.add(user)
            db.session.commit()
            # Send verification email
            token = make_verification_token(user.id)
            verify_url = get_domain() + url_for("verify_email", token=token)
            sent = _send_auth_email(
                user,
                "Verify your RushMail account",
                f"""
                <p>Hi,</p>
                <p>Thanks for signing up to RushMail! Click the button below to verify your email address.
                This link expires in 24 hours.</p>
                <p><a href="{verify_url}" style="background:#f97316;color:#fff;padding:10px 20px;border-radius:6px;text-decoration:none;font-weight:bold">Verify Email</a></p>
                <p>Or paste this link into your browser:<br><a href="{verify_url}">{verify_url}</a></p>
                """,
            )
            if sent:
                flash("Account created! Check your email to verify your address before signing in.", "success")
            else:
                # No app-level or user SMTP available — surface the verify link
                # directly so the user can verify without email delivery.
                flash(
                    f'Account created! Email delivery is not configured, so click this link to verify your account: '
                    f'<a href="{verify_url}">{verify_url}</a>',
                    "success"
                )
            return redirect(url_for("login"))
    return render_template("register.html", password_errors=[])


@app.route("/verify/<token>")
def verify_email(token):
    user_id = verify_verification_token(token)
    if not user_id:
        flash("That verification link is invalid or has expired.", "error")
        return redirect(url_for("login"))
    user = db.session.get(User, user_id)
    if not user:
        flash("Account not found.", "error")
        return redirect(url_for("login"))
    if user.verified:
        flash("Your email is already verified — you can sign in.", "success")
        return redirect(url_for("login"))
    user.verified = True
    db.session.commit()
    flash("Email verified! You can now sign in.", "success")
    return redirect(url_for("login"))


@app.route("/resend-verification")
def resend_verification():
    email = request.args.get("email", "").strip().lower()
    user = User.query.filter_by(email=email).first()
    if user and not user.verified:
        token = make_verification_token(user.id)
        verify_url = get_domain() + url_for("verify_email", token=token)
        _send_auth_email(
            user,
            "Verify your RushMail account",
            f"""
            <p>Hi,</p>
            <p>Here's your new verification link. It expires in 24 hours.</p>
            <p><a href="{verify_url}" style="background:#f97316;color:#fff;padding:10px 20px;border-radius:6px;text-decoration:none;font-weight:bold">Verify Email</a></p>
            <p>Or paste this link:<br><a href="{verify_url}">{verify_url}</a></p>
            """,
        )
    flash("If that address is registered and unverified, we sent a new link. Check your inbox.", "success")
    return redirect(url_for("login"))


@app.route("/forgot-password", methods=["GET", "POST"])
@limiter.limit("5 per hour", methods=["POST"])
def forgot_password():
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        user = User.query.filter_by(email=email).first()
        if user:
            token = make_reset_token(user.id)
            reset_url = get_domain() + url_for("reset_password", token=token)
            _send_auth_email(
                user,
                "Reset your RushMail password",
                f"""
                <p>Hi,</p>
                <p>Someone requested a password reset for your RushMail account.
                Click the button below to set a new password. This link expires in 1 hour.</p>
                <p><a href="{reset_url}" style="background:#f97316;color:#fff;padding:10px 20px;border-radius:6px;text-decoration:none;font-weight:bold">Reset Password</a></p>
                <p>If you didn't request this, you can safely ignore this email.</p>
                <p>Or paste this link:<br><a href="{reset_url}">{reset_url}</a></p>
                """,
            )
        flash("If an account with that email exists, we've sent a reset link. Check your inbox.", "success")
        return redirect(url_for("login"))
    return render_template("forgot_password.html")


@app.route("/reset-password/<token>", methods=["GET", "POST"])
def reset_password(token):
    user_id = verify_reset_token(token)
    if not user_id:
        flash("That reset link is invalid or has expired (links expire after 1 hour).", "error")
        return redirect(url_for("forgot_password"))
    user = db.session.get(User, user_id)
    if not user:
        flash("Account not found.", "error")
        return redirect(url_for("login"))
    if request.method == "POST":
        password  = request.form.get("password", "")
        password2 = request.form.get("password2", "")
        policy_errors = validate_password(password)
        if policy_errors:
            return render_template("reset_password.html", token=token, password_errors=policy_errors)
        elif password != password2:
            return render_template("reset_password.html", token=token, password_errors=[], mismatch=True)
        else:
            user.password_hash = bcrypt.generate_password_hash(password).decode()
            user.verified = True  # also verify if they weren't
            db.session.commit()
            flash("Password updated! You can now sign in.", "success")
            return redirect(url_for("login"))
    return render_template("reset_password.html", token=token, password_errors=[])


@app.route("/logout")
@login_required
def logout():
    logout_user()
    return redirect(url_for("login"))


# ── Dashboard ─────────────────────────────────────────────────────────────────
@app.route("/dashboard")
@login_required
@subscription_required
def dashboard():
    campaigns = (
        Campaign.query
        .filter_by(user_id=current_user.id)
        .order_by(Campaign.created_at.desc())
        .all()
    )
    stats = {
        "campaigns": len(campaigns),
        "sent_ok": sum(c.sent_ok or 0 for c in campaigns),
        "sent_fail": sum(c.sent_fail or 0 for c in campaigns),
    }
    attempted = stats["sent_ok"] + stats["sent_fail"]
    stats["rate"] = round(stats["sent_ok"] / attempted * 100, 1) if attempted else None
    return render_template("dashboard.html", campaigns=campaigns, stats=stats)


# ── Campaign ──────────────────────────────────────────────────────────────────
@app.route("/campaign/new")
@login_required
@subscription_required
def campaign_new():
    return render_template("campaign_new.html")


@app.route("/campaign/<int:campaign_id>")
@login_required
@subscription_required
def campaign_detail(campaign_id):
    campaign = Campaign.query.filter_by(id=campaign_id, user_id=current_user.id).first_or_404()
    recipients = CampaignRecipient.query.filter_by(campaign_id=campaign.id).all()
    return render_template("campaign_detail.html", campaign=campaign, recipients=recipients)


@app.route("/campaign/<int:campaign_id>/resend-failed")
@login_required
@subscription_required
def campaign_resend_failed(campaign_id):
    """Open the campaign wizard pre-filled with this campaign's failed recipients."""
    campaign = Campaign.query.filter_by(id=campaign_id, user_id=current_user.id).first_or_404()
    failed = [
        r.email
        for r in CampaignRecipient.query.filter_by(campaign_id=campaign.id, status="failed").all()
    ]
    if not failed:
        flash("This campaign has no failed recipients to resend to.", "error")
        return redirect(url_for("campaign_detail", campaign_id=campaign.id))
    prefill = {
        "name": f"{campaign.name} (retry)",
        "subject": campaign.subject or "",
        "body": campaign.body or "",
        "emails": failed,
    }
    return render_template("campaign_new.html", prefill=prefill)


@app.route("/campaign/<int:campaign_id>/export.csv")
@login_required
@subscription_required
def campaign_export(campaign_id):
    campaign = Campaign.query.filter_by(id=campaign_id, user_id=current_user.id).first_or_404()
    recipients = CampaignRecipient.query.filter_by(campaign_id=campaign.id).all()
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["email", "status", "sent_at", "error"])
    for r in recipients:
        writer.writerow([
            r.email,
            r.status,
            r.sent_at.isoformat() if r.sent_at else "",
            r.error or "",
        ])
    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", campaign.name or "campaign").strip("_") or "campaign"
    return Response(
        buf.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}_recipients.csv"'},
    )


@app.route("/extract", methods=["POST"])
@login_required
@subscription_required
def extract():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded."}), 400
    f = request.files["file"]
    filename = f.filename.lower()
    text = ""

    try:
        if filename.endswith(".xlsx") or filename.endswith(".xls"):
            wb = openpyxl.load_workbook(f, data_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value:
                            text += str(cell.value) + " "

        elif filename.endswith(".csv"):
            raw = f.read()
            try:
                decoded = raw.decode("utf-8")
            except UnicodeDecodeError:
                decoded = raw.decode("latin-1")
            reader = csv.reader(io.StringIO(decoded))
            for row in reader:
                text += " ".join(row) + " "

        elif filename.endswith(".pdf"):
            import pdfplumber
            with pdfplumber.open(f) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + " "

        elif filename.endswith(".docx"):
            from docx import Document
            doc = Document(f)
            for para in doc.paragraphs:
                text += para.text + " "
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "

        else:
            raw = f.read()
            try:
                text = raw.decode("utf-8")
            except UnicodeDecodeError:
                text = raw.decode("latin-1")

    except Exception as e:
        return jsonify({"error": f"Could not parse file: {e}"}), 400

    emails = extract_emails(text)
    return jsonify({"emails": emails, "total": len(emails)})


def _hostname_is_safe(hostname: str) -> bool:
    """Return False if any resolved IP for hostname is a non-public address."""
    try:
        infos = socket.getaddrinfo(hostname, None)
        if not infos:
            return False
        for info in infos:
            addr = info[4][0]
            ip = ipaddress.ip_address(addr)
            if (
                ip.is_private
                or ip.is_loopback
                or ip.is_link_local
                or ip.is_reserved
                or ip.is_multicast
                or ip.is_unspecified
            ):
                return False
        return True
    except Exception:
        return False


def _is_ssrf_safe(url: str) -> bool:
    """Return False if the URL scheme is not http/https or resolves to a non-public address."""
    try:
        parsed = urlparse(url)
        if parsed.scheme not in ("http", "https"):
            return False
        hostname = parsed.hostname
        if not hostname:
            return False
        return _hostname_is_safe(hostname)
    except Exception:
        return False


def _safe_fetch(url: str, max_redirects: int = 5) -> "req_lib.Response":
    """Fetch a URL, validating SSRF safety at every redirect hop."""
    _HEADERS = {"User-Agent": "Mozilla/5.0"}
    for _ in range(max_redirects + 1):
        if not _is_ssrf_safe(url):
            raise ValueError(f"Blocked: {url} resolves to a private/internal address")
        resp = req_lib.get(url, timeout=10, headers=_HEADERS, allow_redirects=False)
        if resp.is_redirect or resp.status_code in (301, 302, 303, 307, 308):
            location = resp.headers.get("Location", "")
            if not location:
                break
            if location.startswith("/"):
                parsed = urlparse(url)
                url = f"{parsed.scheme}://{parsed.netloc}{location}"
            else:
                url = location
        else:
            resp.raise_for_status()
            return resp
    raise ValueError("Too many redirects or redirect loop detected")


@app.route("/extract-url", methods=["POST"])
@login_required
@subscription_required
def extract_url():
    data = request.get_json()
    url = (data or {}).get("url", "").strip()
    if not url:
        return jsonify({"error": "URL is required."}), 400
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    if not _is_ssrf_safe(url):
        return jsonify({"error": "That URL is not allowed."}), 400
    try:
        resp = _safe_fetch(url)
        soup = BeautifulSoup(resp.text, "html.parser")
        text = soup.get_text(separator=" ")
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": f"Could not fetch URL: {e}"}), 400

    emails = extract_emails(text)
    return jsonify({"emails": emails, "total": len(emails)})


@app.route("/generate", methods=["POST"])
@login_required
@subscription_required
def generate():
    if not OPENAI_API_KEY:
        return jsonify({"error": "OpenAI is not configured. Contact support."}), 503

    data = request.get_json()
    brief = (data or {}).get("brief", "").strip()
    if not brief:
        return jsonify({"error": "Campaign brief is required."}), 400

    try:
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a marketing copywriter. Write a concise, engaging marketing email. "
                        "Return ONLY valid JSON with keys 'subject' (string, max 80 chars) and 'body' (string, plain text, 150-300 words). "
                        "No markdown, no extra keys, no explanations."
                    )
                },
                {
                    "role": "user",
                    "content": f"Campaign brief: {brief}"
                }
            ],
            response_format={"type": "json_object"},
            max_tokens=600,
        )
        result = json.loads(response.choices[0].message.content)
        if "subject" not in result or "body" not in result:
            raise ValueError("Missing keys in AI response")
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": f"AI generation failed: {e}"}), 500


@app.route("/send-bulk", methods=["POST"])
@login_required
@subscription_required
def send_bulk():
    # Accept multipart/form-data (with optional file attachments)
    emails_raw = request.form.get("emails", "[]")
    subject = request.form.get("subject", "").strip()
    body = request.form.get("body", "").strip()
    campaign_name = (request.form.get("name", "Campaign") or "Campaign").strip()

    try:
        emails = json.loads(emails_raw)
    except Exception:
        return jsonify({"error": "Invalid emails payload."}), 400

    body_text = re.sub(r"<[^>]+>", "", body).strip()
    if not emails or not subject or not body_text:
        return jsonify({"error": "emails, subject, and body are required."}), 400

    if not current_user.smtp_host:
        return jsonify({"error": "SMTP is not configured. Go to Settings."}), 400

    # Read uploaded attachments into memory
    attachment_files = request.files.getlist("attachments")
    attachments = []
    for att_file in attachment_files:
        if att_file and att_file.filename:
            att_data = att_file.read()
            attachments.append((att_file.filename, att_data, att_file.mimetype or "application/octet-stream"))

    smtp_host = current_user.smtp_host
    smtp_port = current_user.smtp_port
    smtp_user = current_user.smtp_user
    smtp_pass = decrypt_password(current_user.smtp_pass_enc)
    use_tls = current_user.smtp_use_tls
    from_addr = current_user.smtp_from or smtp_user

    # Create campaign record
    campaign = Campaign(
        user_id=current_user.id,
        name=campaign_name,
        subject=subject,
        body=body,
        total=len(emails),
    )
    db.session.add(campaign)
    db.session.flush()

    recipient_rows = []
    for email in emails:
        r = CampaignRecipient(campaign_id=campaign.id, email=email)
        db.session.add(r)
        recipient_rows.append(r)
    db.session.flush()

    # Capture all data we need from ORM objects as plain Python values
    # BEFORE commit() expires them — avoids DetachedInstanceError in the generator.
    campaign_id = campaign.id
    recipient_data = [{"id": r.id, "email": r.email} for r in recipient_rows]
    db.session.commit()

    def stream():
        ok = 0
        fail = 0
        server = None
        try:
            if use_tls:
                server = smtplib.SMTP(smtp_host, smtp_port, timeout=10)
                server.ehlo()
                server.starttls()
            else:
                server = smtplib.SMTP_SSL(smtp_host, smtp_port, timeout=10)
            server.login(smtp_user, smtp_pass)
        except Exception as e:
            err_str = str(e)
            with app.app_context():
                for rd in recipient_data:
                    CampaignRecipient.query.filter_by(id=rd["id"]).update(
                        {"status": "failed", "error": err_str}
                    )
                    fail += 1
                    yield json.dumps({"email": rd["email"], "status": "failed", "error": err_str}) + "\n"
                Campaign.query.filter_by(id=campaign_id).update(
                    {"sent_ok": 0, "sent_fail": fail, "status": "completed"}
                )
                db.session.commit()
            yield json.dumps({"done": True, "ok": 0, "fail": fail, "campaign_id": campaign_id}) + "\n"
            return

        for rd in recipient_data:
            r_email = rd["email"]
            r_id = rd["id"]
            try:
                msg = MIMEMultipart("mixed")
                msg["From"] = from_addr
                msg["To"] = r_email
                msg["Subject"] = subject
                msg.attach(MIMEText(body, "html"))

                for att_name, att_data, att_mime in attachments:
                    maintype, subtype = att_mime.split("/", 1) if "/" in att_mime else ("application", "octet-stream")
                    part = MIMEBase(maintype, subtype)
                    part.set_payload(att_data)
                    encoders.encode_base64(part)
                    part.add_header("Content-Disposition", "attachment", filename=att_name)
                    msg.attach(part)

                server.sendmail(from_addr, r_email, msg.as_string())
                with app.app_context():
                    CampaignRecipient.query.filter_by(id=r_id).update(
                        {"status": "sent", "sent_at": datetime.utcnow()}
                    )
                    db.session.commit()
                ok += 1
                yield json.dumps({"email": r_email, "status": "sent"}) + "\n"
            except Exception as e:
                err_str = str(e)
                with app.app_context():
                    CampaignRecipient.query.filter_by(id=r_id).update(
                        {"status": "failed", "error": err_str}
                    )
                    db.session.commit()
                fail += 1
                yield json.dumps({"email": r_email, "status": "failed", "error": err_str}) + "\n"

        try:
            server.quit()
        except Exception:
            pass

        with app.app_context():
            Campaign.query.filter_by(id=campaign_id).update(
                {"sent_ok": ok, "sent_fail": fail, "status": "completed"}
            )
            db.session.commit()
        yield json.dumps({"done": True, "ok": ok, "fail": fail, "campaign_id": campaign_id}) + "\n"

    return Response(stream_with_context(stream()), mimetype="application/x-ndjson")


# ── Settings ──────────────────────────────────────────────────────────────────
@app.route("/settings", methods=["GET", "POST"])
@login_required
@subscription_required
def settings():
    if request.method == "POST":
        current_user.smtp_host = request.form.get("smtp_host", "").strip() or None
        current_user.smtp_port = int(request.form.get("smtp_port", 587) or 587)
        current_user.smtp_user = request.form.get("smtp_user", "").strip() or None
        current_user.smtp_from = request.form.get("smtp_from", "").strip() or None
        current_user.smtp_use_tls = "smtp_use_tls" in request.form
        new_pass = request.form.get("smtp_pass", "").strip()
        if new_pass:
            current_user.smtp_pass_enc = encrypt_password(new_pass)
        db.session.commit()
        flash("Settings saved.", "success")
        return redirect(url_for("settings"))
    return render_template("settings.html", stripe_enabled=STRIPE_ENABLED)


@app.route("/settings/test-email", methods=["POST"])
@login_required
@subscription_required
@limiter.limit("5 per minute")
def settings_test_email():
    """Send a test email to the user's own address using their saved SMTP settings."""
    if not (current_user.smtp_host and current_user.smtp_user and current_user.smtp_pass_enc):
        return jsonify({"error": "Save your SMTP settings first, then send a test."}), 400
    try:
        smtp_pass = decrypt_password(current_user.smtp_pass_enc)
    except Exception:
        return jsonify({"error": "Stored SMTP password could not be decrypted. Re-enter and save it."}), 400

    from_addr = current_user.smtp_from or current_user.smtp_user
    msg = MIMEText(
        "This is a test email from RushMail.\n\n"
        "If you're reading this, your SMTP settings are working and you're ready to send campaigns.",
        "plain",
    )
    msg["Subject"] = "RushMail SMTP test"
    msg["From"] = from_addr
    to_addr = current_user.smtp_from or current_user.smtp_user
    msg["To"] = to_addr

    try:
        if current_user.smtp_use_tls:
            server = smtplib.SMTP(current_user.smtp_host, current_user.smtp_port, timeout=15)
            server.ehlo()
            server.starttls()
        else:
            server = smtplib.SMTP_SSL(current_user.smtp_host, current_user.smtp_port, timeout=15)
        server.login(current_user.smtp_user, smtp_pass)
        server.sendmail(from_addr, [to_addr], msg.as_string())
        server.quit()
    except Exception as e:
        return jsonify({"error": f"{type(e).__name__}: {e}"}), 400

    return jsonify({"ok": True, "message": f"Test email sent to {to_addr} — check your inbox."})


# ── Stripe / Pricing ──────────────────────────────────────────────────────────
@app.route("/pricing")
@login_required
def pricing():
    return render_template("pricing.html", stripe_enabled=STRIPE_ENABLED)


@app.route("/subscribe")
@login_required
def subscribe():
    if not STRIPE_ENABLED:
        flash("Payments are not configured yet.", "error")
        return redirect(url_for("pricing"))
    if not STRIPE_PRO_PRICE_ID:
        flash("Pro plan price not configured. Contact support.", "error")
        return redirect(url_for("pricing"))

    try:
        if not current_user.stripe_customer_id:
            customer = stripe.Customer.create(email=current_user.email)
            current_user.stripe_customer_id = customer.id
            db.session.commit()

        domain = get_domain()
        session = stripe.checkout.Session.create(
            customer=current_user.stripe_customer_id,
            payment_method_types=["card"],
            line_items=[{"price": STRIPE_PRO_PRICE_ID, "quantity": 1}],
            mode="subscription",
            success_url=f"{domain}/subscribe/success?session_id={{CHECKOUT_SESSION_ID}}",
            cancel_url=f"{domain}/pricing",
        )
        return redirect(session.url, code=303)
    except stripe.error.StripeError as e:
        flash(f"Stripe error: {e.user_message}", "error")
        return redirect(url_for("pricing"))


@app.route("/subscribe/success")
@login_required
def subscribe_success():
    session_id = request.args.get("session_id")
    if session_id and STRIPE_ENABLED:
        try:
            session = stripe.checkout.Session.retrieve(session_id)
            if (
                session.customer
                and current_user.stripe_customer_id
                and session.customer == current_user.stripe_customer_id
                and session.payment_status in ("paid", "no_payment_required")
                and session.subscription
            ):
                current_user.stripe_subscription_id = session.subscription
                current_user.plan = "pro"
                db.session.commit()
        except Exception:
            pass
    flash("Welcome to RushMail Pro! Your account is now active.", "success")
    return redirect(url_for("dashboard"))


@app.route("/billing-portal")
@login_required
def billing_portal():
    if not STRIPE_ENABLED or not current_user.stripe_customer_id:
        flash("Billing portal is not available.", "error")
        return redirect(url_for("settings"))
    try:
        domain = get_domain()
        portal = stripe.billing_portal.Session.create(
            customer=current_user.stripe_customer_id,
            return_url=f"{domain}/settings",
        )
        return redirect(portal.url, code=303)
    except stripe.error.StripeError as e:
        flash(f"Could not open billing portal: {e.user_message}", "error")
        return redirect(url_for("settings"))


@app.route("/webhook", methods=["POST"])
@csrf.exempt
def stripe_webhook():
    payload = request.data
    sig_header = request.headers.get("Stripe-Signature", "")

    if not STRIPE_ENABLED or not STRIPE_WEBHOOK_SECRET:
        return jsonify({"error": "Webhooks not configured"}), 400

    try:
        event = stripe.Webhook.construct_event(payload, sig_header, STRIPE_WEBHOOK_SECRET)
    except (ValueError, stripe.error.SignatureVerificationError):
        return jsonify({"error": "Invalid payload"}), 400

    if event["type"] == "customer.subscription.updated":
        sub = event["data"]["object"]
        user = User.query.filter_by(stripe_customer_id=sub["customer"]).first()
        if user:
            user.stripe_subscription_id = sub["id"]
            user.plan = "pro" if sub["status"] in ("active", "trialing") else "free"
            db.session.commit()

    elif event["type"] == "customer.subscription.deleted":
        sub = event["data"]["object"]
        user = User.query.filter_by(stripe_customer_id=sub["customer"]).first()
        if user:
            user.plan = "free"
            user.stripe_subscription_id = None
            db.session.commit()

    return jsonify({"received": True})


# ── Scheduled campaigns ───────────────────────────────────────────────────────
def _fire_scheduled_campaign(sc_id: int, smtp_cfg: dict):
    """Send one scheduled campaign run. Called from the scheduler background thread."""
    with app.app_context():
        sc = ScheduledCampaign.query.get(sc_id)
        if not sc:
            return
        emails = sc.emails
        subject = sc.subject
        body = sc.body
        from_addr = smtp_cfg["from"] or smtp_cfg["user"]
        ok = fail = 0
        server = None
        try:
            if smtp_cfg["use_tls"]:
                server = smtplib.SMTP(smtp_cfg["host"], smtp_cfg["port"], timeout=20)
                server.ehlo()
                server.starttls()
            else:
                server = smtplib.SMTP_SSL(smtp_cfg["host"], smtp_cfg["port"], timeout=20)
            server.login(smtp_cfg["user"], smtp_cfg["pass"])
            for addr in emails:
                try:
                    msg = MIMEMultipart("mixed")
                    msg["From"] = from_addr
                    msg["To"] = addr
                    msg["Subject"] = subject
                    msg.attach(MIMEText(body, "html"))
                    server.sendmail(from_addr, addr, msg.as_string())
                    ok += 1
                except Exception:
                    fail += 1
            server.quit()
        except Exception:
            fail += len(emails)
        # Record as a Campaign for history
        c = Campaign(
            user_id=sc.user_id,
            name=f"[Scheduled] {sc.name}",
            subject=subject,
            body=body,
            total=len(emails),
            sent_ok=ok,
            sent_fail=fail,
            status="completed",
        )
        db.session.add(c)
        db.session.commit()


FREQUENCY_CHOICES = ("once", "daily", "weekly", "monthly")


def _next_occurrence(dt: datetime, frequency: str) -> datetime:
    """Return the next run time after `dt` for the given frequency."""
    if frequency == "daily":
        return dt + timedelta(days=1)
    if frequency == "monthly":
        # Add one calendar month, clamping the day if the target month is shorter.
        month = dt.month + 1
        year = dt.year + (month - 1) // 12
        month = (month - 1) % 12 + 1
        import calendar
        day = min(dt.day, calendar.monthrange(year, month)[1])
        return dt.replace(year=year, month=month, day=day)
    # "weekly" (and any unrecognized value) falls back to weekly.
    return dt + timedelta(weeks=1)


def run_scheduled_campaigns():
    """Checked every minute by APScheduler. Fires any due campaigns."""
    with app.app_context():
        now = datetime.utcnow()
        due = ScheduledCampaign.query.filter(
            ScheduledCampaign.active == True,
            ScheduledCampaign.next_run_at <= now,
        ).all()
        for sc in due:
            old_next = sc.next_run_at
            is_one_off = sc.frequency == "once"
            # Optimistic lock: only one worker proceeds if next_run_at still matches
            updated = db.session.execute(
                db.update(ScheduledCampaign)
                .where(
                    ScheduledCampaign.id == sc.id,
                    ScheduledCampaign.next_run_at == old_next,
                )
                .values(
                    next_run_at=old_next if is_one_off else _next_occurrence(old_next, sc.frequency),
                    last_run_at=now,
                    active=(False if is_one_off else ScheduledCampaign.active),
                )
            )
            db.session.commit()
            if updated.rowcount == 0:
                continue  # Another worker already claimed this run
            user = User.query.get(sc.user_id)
            if not user or not user.smtp_host or not user.smtp_pass_enc:
                continue
            smtp_cfg = {
                "host": user.smtp_host,
                "port": user.smtp_port,
                "user": user.smtp_user,
                "pass": decrypt_password(user.smtp_pass_enc),
                "use_tls": user.smtp_use_tls,
                "from": user.smtp_from,
            }
            _fire_scheduled_campaign(sc.id, smtp_cfg)


def _validate_schedule_input(name, subject, body, emails_list, first_run, frequency):
    """Shared validation for creating a ScheduledCampaign from any entry point."""
    errors = []
    if not name:
        errors.append("Campaign name is required.")
    if not subject:
        errors.append("Subject is required.")
    if not body:
        errors.append("Body is required.")
    if not emails_list:
        errors.append("At least one email address is required.")
    if frequency not in FREQUENCY_CHOICES:
        errors.append("Invalid frequency.")

    next_run_at = None
    if not first_run:
        errors.append("First send date/time is required.")
    else:
        try:
            next_run_at = datetime.strptime(first_run, "%Y-%m-%dT%H:%M")
        except ValueError:
            errors.append("Invalid date/time format.")

    return errors, next_run_at


@app.route("/scheduled", methods=["GET", "POST"])
@login_required
@subscription_required
def scheduled():
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        subject = request.form.get("subject", "").strip()
        body = request.form.get("body", "").strip()
        emails_raw = request.form.get("emails", "").strip()
        first_run = request.form.get("first_run", "").strip()
        frequency = request.form.get("frequency", "weekly").strip().lower()

        emails_list = [e.strip().lower() for e in emails_raw.splitlines() if e.strip()]
        emails_list = list(dict.fromkeys(emails_list))  # deduplicate

        errors, next_run_at = _validate_schedule_input(name, subject, body, emails_list, first_run, frequency)

        if errors:
            for e in errors:
                flash(e, "error")
        else:
            sc = ScheduledCampaign(
                user_id=current_user.id,
                name=name,
                subject=subject,
                body=body,
                emails_json=json.dumps(emails_list),
                next_run_at=next_run_at,
                frequency=frequency,
            )
            db.session.add(sc)
            db.session.commit()
            flash(f'{frequency.capitalize()} schedule "{name}" created — first send on {next_run_at.strftime("%b %d, %Y at %H:%M")} UTC.', "success")
        return redirect(url_for("scheduled"))

    schedules = ScheduledCampaign.query.filter_by(user_id=current_user.id).order_by(ScheduledCampaign.created_at.desc()).all()
    return render_template("scheduled.html", schedules=schedules)


@app.route("/campaign/schedule", methods=["POST"])
@login_required
@subscription_required
def campaign_schedule():
    """Create a scheduled campaign directly from the New Campaign wizard."""
    data = request.get_json(silent=True) or {}
    name = (data.get("name") or "").strip()
    subject = (data.get("subject") or "").strip()
    body = (data.get("body") or "").strip()
    first_run = (data.get("first_run") or "").strip()
    frequency = (data.get("frequency") or "weekly").strip().lower()
    emails_list = data.get("emails") or []
    emails_list = [str(e).strip().lower() for e in emails_list if str(e).strip()]
    emails_list = list(dict.fromkeys(emails_list))

    errors, next_run_at = _validate_schedule_input(name, subject, body, emails_list, first_run, frequency)
    if errors:
        return jsonify({"error": " ".join(errors)}), 400

    sc = ScheduledCampaign(
        user_id=current_user.id,
        name=name,
        subject=subject,
        body=body,
        emails_json=json.dumps(emails_list),
        next_run_at=next_run_at,
        frequency=frequency,
    )
    db.session.add(sc)
    db.session.commit()
    return jsonify({
        "ok": True,
        "message": f'{frequency.capitalize()} schedule "{name}" created — first send on {next_run_at.strftime("%b %d, %Y at %H:%M")} UTC.',
    })


@app.route("/scheduled/<int:sc_id>/toggle", methods=["POST"])
@login_required
@subscription_required
def scheduled_toggle(sc_id):
    sc = ScheduledCampaign.query.filter_by(id=sc_id, user_id=current_user.id).first_or_404()
    sc.active = not sc.active
    db.session.commit()
    state = "resumed" if sc.active else "paused"
    flash(f'Schedule "{sc.name}" {state}.', "success")
    return redirect(url_for("scheduled"))


@app.route("/scheduled/<int:sc_id>/delete", methods=["POST"])
@login_required
@subscription_required
def scheduled_delete(sc_id):
    sc = ScheduledCampaign.query.filter_by(id=sc_id, user_id=current_user.id).first_or_404()
    db.session.delete(sc)
    db.session.commit()
    flash(f'Schedule "{sc.name}" deleted.', "success")
    return redirect(url_for("scheduled"))


# ── Help ──────────────────────────────────────────────────────────────────────
@app.route("/help")
@login_required
def help_page():
    return render_template("help.html")


# ── Tutorial ──────────────────────────────────────────────────────────────────
@app.route("/tutorial")
@login_required
def tutorial():
    return render_template("tutorial.html")


# ── Init ──────────────────────────────────────────────────────────────────────
with app.app_context():
    db.create_all()
    # Lightweight migrations: db.create_all() won't add columns to existing tables.
    try:
        from sqlalchemy import inspect, text
        inspector = inspect(db.engine)
        if "scheduled_campaigns" in inspector.get_table_names():
            cols = [c["name"] for c in inspector.get_columns("scheduled_campaigns")]
            if "frequency" not in cols:
                db.session.execute(text(
                    "ALTER TABLE scheduled_campaigns ADD COLUMN frequency VARCHAR(20) NOT NULL DEFAULT 'weekly'"
                ))
                db.session.commit()
        if "users" in inspector.get_table_names():
            user_cols = [c["name"] for c in inspector.get_columns("users")]
            if "verified" not in user_cols:
                db.session.execute(text(
                    "ALTER TABLE users ADD COLUMN verified BOOLEAN NOT NULL DEFAULT TRUE"
                ))
                db.session.commit()
    except Exception:
        db.session.rollback()

# Start background scheduler (only in the reloader child in dev; always in prod)
if os.environ.get("WERKZEUG_RUN_MAIN") == "true" or not app.debug:
    _scheduler = BackgroundScheduler(daemon=True)
    _scheduler.add_job(run_scheduled_campaigns, "interval", minutes=1, max_instances=1)
    _scheduler.start()

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
